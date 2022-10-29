#world cloud are presented in word file and we have stored them in folder
import pandas as pd
import numpy as npy
import matplotlib.pyplot as pPlot
import docx
from docx import Document
import operator
import re
import itertools
from collections import Counter
import nltk
from nltk.stem import WordNetLemmatizer
from nltk.stem.porter import PorterStemmer
from wordcloud import WordCloud, STOPWORDS

doc = Document('cs1.docx')

n=len(doc.tables)
for i in range(0,n):
   doc.tables[0]._element.getparent().remove(doc.tables[0]._element)     

#docs after removing table is stored as t1.docs

doc.save('t1.docx')

t2 =Document('t1.docx')
doc_1 = docx.Document()
for para in t2.paragraphs :
   new_text = para.text
   temp = re.sub(r'Fig. [0-9][0-9].[0-9]',' ' , new_text)
   temp1 = re.sub(r'Fig. [0-9].[0-9]',' ' , temp)
   temp2=re.sub(r'Fig. [0-9].[0-9][0-9]',' ' , temp1)
   temp3=re.sub(r'Fig. [0-9][0-9].[0-9][0-9]',' ' , temp2)
   temp4=re.sub(r'Table [0-9].[0-9]',' ' , temp3)
   temp5=re.sub(r'Table [0-9][0-9].[0-9]',' ' , temp4)
   temp6=re.sub(r'Sects. [0-9].[0-9]',' ' , temp5)
   temp7=re.sub(r'Sects. [0-9][0-9].[0-9]',' ' , temp6)
   temp8=re.sub(r'[0-9][0-9].[0-9]',' ' , temp7)
   temp9=re.sub(r'[0-9].[0-9]',' ' , temp8)
   temp10=re.sub(r'Chap. [0-9]',' ' , temp9)
   temp11 = re.sub(r"[^a-zA-Z0-9]", " ", temp10)
   temp12=re.sub(r'Chap. [0-9][0-9]',' ' , temp11)
   if len(temp12)>20 :
      doc_1.add_paragraph(temp12)

doc_1.save('t3.docx')

#docs after removing images from t1 is stored as t3.docs

t4=Document('t3.docx')

# tockenisation
tokens = []
for para in t4.paragraphs:
   nltk_tokens = nltk.word_tokenize(para.text,"english",False)
   tokens=tokens+nltk_tokens


tokens_final=[]
for ele in tokens : 
   if ele ==',' or ele=='.' or ele=='(' or ele==')' or ele=='=' or ele==';' or ele=='1' or ele== '+' or ele=='2' or ele=='[' or ele==']' :
      tokens.remove(ele)

#lemmatisation

wordnet_lemmatizer = WordNetLemmatizer()
lema=[]
for w in tokens:
   lema.append(wordnet_lemmatizer.lemmatize(w))


# steming
# nltk.download('wordnet')
stem=[]
porter_stemmer = PorterStemmer()
for w in lema:
   stem.append(porter_stemmer.stem(w))

for ele1 in stem :
   if ele1==':' :
      stem.remove(ele1)



#frequency analyses
word_freq = {}
for i in range(len(stem)):
    if stem[i] in word_freq.keys():
        word_freq[stem[i]]+=1
    else:
        word_freq[stem[i]]=1

word_freq_sorted = dict( sorted(word_freq.items(), key=operator.itemgetter(1),reverse=True))


n3=[]
n4=[]
out = dict(itertools.islice(word_freq_sorted.items(), 20))
for k in out.keys():
   n3.append(k)                    
   n4.append(out[k])

pPlot.bar(n3,n4)
pPlot.xticks(rotation=90)             # graph is saved as graph1.png
pPlot.show()

n1=[]
n2=[]
out = dict(itertools.islice(word_freq_sorted.items(), 20))
for k in out.keys():
   n1.append(len(k))                    
   n2.append(out[k])


pPlot.bar(n1,n2)
pPlot.xticks(rotation=90)             # graph is saved as graph.png
pPlot.show()

#forming wordcoud1
dataset = open("stemm.txt", "r",encoding='utf-8').read()

def create_word_cloud(string):
   # maskArray = npy.array(Image.open("cloud.png"))
   cloud = WordCloud(background_color = "white", max_words = 200,  stopwords = set(STOPWORDS))
   cloud.generate(string)
   cloud.to_file("wordCloud.png")                  #first wordcloud is formed named as wordcloud.png

dataset = dataset.lower()

create_word_cloud(dataset)


#removing stop words
after_remove= []
from nltk.corpus import stopwords
stop = set(stopwords.words('english'))
for word in stem: 
    if word not in stop:
      after_remove.append(word)

for ele1 in after_remove : 
   if ele1 =='e' or ele1=='.' or ele1=='(' or ele1==')' or ele1=='=' or ele1==';' or ele1=='1' or ele1== '+' or ele1=='2' or ele1=='[' or ele1==']' or ele1=="x" or ele1 =='p':
      after_remove.remove(ele1)


word_freq_1 = {}
for i in range(len(after_remove)):
    if after_remove[i] in word_freq_1.keys():
        word_freq_1[after_remove[i]]+=1
    else:
        word_freq_1[after_remove[i]]=1

word_freq_sorted_1 = dict( sorted(word_freq_1.items(), key=operator.itemgetter(1),reverse=True))

n5=[]
n6=[]
out = dict(itertools.islice(word_freq_sorted_1.items(), 20))
for k in out.keys():
   n5.append(k)                    
   n6.append(out[k])

pPlot.bar(n5,n6)
pPlot.xticks(rotation=90)             # graph is saved as graph1.png
pPlot.show()


def create_word_cloud1(string):
   # maskArray = npy.array(Image.open("cloud.png"))
   cloud = WordCloud(background_color = "white", max_words = 200,  stopwords = set(STOPWORDS))
   cloud.generate(string)
   cloud.to_file("wordCloud1.png")                 #second wordcloud is formed named as wordcloud1.png

dataset1 = open("remove.txt", "r",encoding='utf-8').read()
dataset1 = dataset1.lower()
create_word_cloud1(dataset1)

pos = nltk.pos_tag(stem)

the_count = Counter(tag for _, tag in pos)

def FrequencyPlot(distribution):
  #plt.rcParams["figure.autolayout"] = True
  pPlot.rcParams["figure.figsize"] = [15, 3.50]
  pPlot.bar(distribution.keys(),distribution.values())
  pPlot.xticks(rotation=90)
  pPlot.show()

FrequencyPlot(the_count)


print(the_count)



